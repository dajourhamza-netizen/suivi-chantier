import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from datetime import datetime
import os
import re
import io

from docx import Document
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

# Configuration de la page
st.set_page_config(page_title="Suivi de Chantier", page_icon="🏗️", layout="wide")

FILE_PATH = "suivi .xlsx"
COL_PARTIE = "PARTIE D'OUVRAGE"

COLUMNS_TEMPLATE = [
    "DATE", "TITRE DE LA NATURE DES TRAVAUX", COL_PARTIE, 
    "SITUATION", "ACTIVITÉ RÉALISÉE", "ÉSSAI/ CONTRÔLE RÉALISÉE", 
    "RÉFÉRENCE DE PROCÉDURE", "PIÈCES JOINTES"
]

LIAISONS = {
    "ARASE DE PST": {"procedure": "TER-PEX-05-00", "pieces": "* Fiche de suivi de la PST\n* Fiche de réception topographique\n* PVs laboratoire"},
    "ARASE DE TERRASSEMENT": {"procedure": "TER-PEX-03-00", "pieces": "* Fiche de contrôle des déblais\n* Fiche de réception topographique\n* PVs laboratoire"},
    "ASSISE DE REMBLAIS PURGE": {"procedure": "TER-PEX-04-00", "pieces": "* Fiche de réception de l'assise des remblais\n* Fiche de réception topographique\n* Fiche d'identification de la purge\n* PVs laboratoire"},
    "ASSISE DE REMBLAIS": {"procedure": "TER-PEX-04-00", "pieces": "* Fiche de réception de l'assise des remblais\n* Fiche de réception topographique\n* PVs laboratoire"},
    "ASSISE DE REMBLAIS CDF": {"procedure": "TER-PEX-04-00", "pieces": "* Fiche de réception de l'assise des remblais\n* Fiche de réception topographique\n* PVs laboratoire"},
    "ASSISE DE REMBLAIS CONTIGUS": {"procedure": "OVA-PEX-16-00", "pieces": "* Fiche de suivi des remblais contigus\n* Fiche de contrôle des remblais contigus\n* PVs laboratoire\n* Fiche de réception topographique"},
    "ASSISE DE REMBLAI DE FOUILLE": {"procedure": "OVA-PEX-04-00", "pieces": "* Fiche de suivi et de contrôle des fouilles et remblaiement de fouilles\n* PVs laboratoire"},
    "ASSISE DE REMBLAIS RENFORCE": {"procedure": "TER-PEX-13-00", "pieces": "* PV Manifold\n* PVs laboratoire\n* Fiche de réception topographique\n* Fiche de réception assise remblai renforcé"},
    "ASSISE DRAINANTE": {"procedure": "TER-PEX-13-00", "pieces": "* Fiche de réception topographique\n* PVs laboratoire\n* Fiche de contrôle de l'assise drainante"},
    "COUCHE DE FORME": {"procedure": "TER-PEX-09-00", "pieces": "* Fiche de suivi et de contrôle de la CDF\n* Fiche de réception topographique\n* PVs laboratoire"},
    "DÉCAPAGE": {"procedure": "TER-PEX-02-00", "pieces": "* Fiche de suivi et de contrôle du décapage\n* Fiche des sections à décaper\n* Fiche de réception topographique"},
    "DEGAGEMENT D'EMPRISE": {"procedure": "TER-PEX-01-00", "pieces": "* Fiche de suivi et de contrôle du dégagement des emprises\n* Fiche de réception topographique\n* Constat dégagement d'emprise"},
    "REMBLAIS": {"procedure": "TER-PEX-04-00", "pieces": "* Fiche de suivi et de contrôle des remblais\n* PVs laboratoire"},
    "REMBLAIS CDF": {"procedure": "TER-PEX-04-00", "pieces": "* Fiche de suivi et de contrôle des remblais\n* PVs laboratoire"},
    "REMBLAIS CONTIGUS": {"procedure": "OVA-PEX-16-00", "pieces": "* Fiche de suivi des remblais contigus\n* Fiche de contrôle des remblais contigus\n* PVs laboratoire\n* Fiche de réception topographique"},
    "REMBLAIS DE FOUILLE": {"procedure": "OVA-PEX-04-00", "pieces": "* Fiche de suivi et de contrôle des fouilles et remblaiement de fouilles\n* PVs laboratoire"},
    "REMBLAIS DE FOUILLS CDF": {"procedure": "OVA-PEX-04-00", "pieces": "* Fiche de suivi et de contrôle des fouilles et remblaiement de fouilles\n* PVs laboratoire"},
    "REMBLAIS RENFORCE": {"procedure": "TER-PEX-13-00", "pieces": "* Fiche de suivi des remblais renforcé\n* Fiche de contrôle des armatures Geostrap\n* Fiche de réception de pose des ecailles\n* PVs laboratoire"},
    "REMBLAIS PST": {"procedure": "TER-PEX-05-00", "pieces": "* Fiche de suivi et de contrôle des remblais PST\n* PVs laboratoire"}
}

# ==========================================
# FONCTION 1 : GENERER LA PAGE D'IMPRESSION GLOBAL DU TABLEAU
# ==========================================
def generate_printable_html(df_data, chantier_name):
    now_str = datetime.today().strftime('%d/%m/%Y à %H:%M')
    rows_html = ""
    for idx, row in df_data.iterrows():
        date_val = str(row.get('DATE', '') or '')
        nature_val = str(row.get('TITRE DE LA NATURE DES TRAVAUX', '') or '')
        partie_val = str(row.get(COL_PARTIE, '') or '')
        situ_val = str(row.get('SITUATION', '') or '')
        act_val = str(row.get('ACTIVITÉ RÉALISÉE', '') or '')
        essai_val = str(row.get('ÉSSAI/ CONTRÔLE RÉALISÉE', '') or '')
        proc_val = str(row.get('RÉFÉRENCE DE PROCÉDURE', '') or '')
        
        rows_html += f"""
        <tr>
            <td style="text-align:center; font-weight:bold;">{idx + 1}</td>
            <td style="white-space:nowrap;">{date_val}</td>
            <td><b>{nature_val}</b></td>
            <td>{partie_val}</td>
            <td>{situ_val}</td>
            <td>{act_val}</td>
            <td>{essai_val}</td>
            <td>{proc_val}</td>
        </tr>
        """

    return f"""
    <!DOCTYPE html>
    <html lang="fr">
    <head>
        <meta charset="UTF-8">
        <title>Impression Tableau - {chantier_name}</title>
        <style>
            @page {{ size: A4 landscape; margin: 10mm; }}
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 10px; font-size: 11px; color: #1e293b; background-color: #ffffff; }}
            .header-banner {{ display: flex; justify-content: space-between; align-items: center; border-bottom: 2px solid #1e3a8a; padding-bottom: 8px; margin-bottom: 15px; }}
            .title {{ font-size: 18px; font-weight: bold; color: #1e3a8a; text-transform: uppercase; }}
            .subtitle {{ font-size: 12px; color: #475569; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 5px; }}
            th, td {{ border: 1px solid #cbd5e1; padding: 6px 8px; text-align: left; vertical-align: top; }}
            th {{ background-color: #1e293b; color: #ffffff; font-size: 10px; font-weight: bold; text-transform: uppercase; letter-spacing: 0.5px; }}
            tr:nth-child(even) {{ background-color: #f8fafc; }}
            .btn-container {{ margin-bottom: 15px; text-align: right; }}
            .print-btn {{ background-color: #0284c7; color: white; border: none; padding: 10px 22px; font-size: 14px; font-weight: bold; border-radius: 6px; cursor: pointer; box-shadow: 0 2px 4px rgba(0,0,0,0.15); }}
            @media print {{ .no-print {{ display: none !important; }} body {{ margin: 0; }} }}
        </style>
    </head>
    <body>
        <div class="no-print btn-container">
            <button class="print-btn" onclick="window.print()">🖨️ Lancer l'Impression / Enregistrer en PDF</button>
        </div>
        <div class="header-banner">
            <div>
                <div class="title">🏗️ RAPPORT DE SUIVI DE CHANTIER</div>
                <div class="subtitle">Chantier : <b>{chantier_name}</b></div>
            </div>
            <div style="text-align:right;">
                <div class="subtitle">Édité le : <b>{now_str}</b></div>
                <div class="subtitle">Nombre de lignes : <b>{len(df_data)}</b></div>
            </div>
        </div>
        <table>
            <thead>
                <tr>
                    <th style="width: 30px;">N°</th>
                    <th style="width: 80px;">Date</th>
                    <th>Nature des Travaux</th>
                    <th>Partie d'Ouvrage</th>
                    <th>Situation / PK</th>
                    <th>Activité Réalisée</th>
                    <th>Essai / Contrôle</th>
                    <th>Procédure</th>
                </tr>
            </thead>
            <tbody>
                {rows_html}
            </tbody>
        </table>
    </body>
    </html>
    """

# ==========================================
# FONCTION 2 : GENERER LE DOCUMENT FICHE INDIVIDUELLE PDF
# ==========================================
def generate_single_document_html(row_data, chantier_name):
    date_val = str(row_data.get('DATE', '') or '-')
    nature_val = str(row_data.get('TITRE DE LA NATURE DES TRAVAUX', '') or '-')
    partie_val = str(row_data.get(COL_PARTIE, '') or '-')
    situ_val = str(row_data.get('SITUATION', '') or '-')
    act_val = str(row_data.get('ACTIVITÉ RÉALISÉE', '') or '-')
    essai_val = str(row_data.get('ÉSSAI/ CONTRÔLE RÉALISÉE', '') or '-')
    proc_val = str(row_data.get('RÉFÉRENCE DE PROCÉDURE', '') or '-')
    pieces_val = str(row_data.get('PIÈCES JOINTES', '') or '-').replace('\n', '<br>')

    return f"""
    <!DOCTYPE html>
    <html lang="fr">
    <head>
        <meta charset="UTF-8">
        <title>{nature_val}_{partie_val} _.pdf</title>
        <style>
            @page {{ size: A4 portrait; margin: 15mm; }}
            body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 13px; color: #0f172a; margin: 10px; background-color: #fff; }}
            .header-box {{ border: 2px solid #1e3a8a; padding: 15px; text-align: center; margin-bottom: 20px; background-color: #f8fafc; border-radius: 6px; }}
            .header-title {{ font-size: 20px; font-weight: bold; color: #1e3a8a; margin-bottom: 5px; text-transform: uppercase; }}
            .header-sub {{ font-size: 13px; color: #475569; font-weight: 600; }}
            
            .doc-table {{ width: 100%; border-collapse: collapse; margin-top: 15px; }}
            .doc-table th, .doc-table td {{ border: 1px solid #94a3b8; padding: 12px 14px; text-align: left; vertical-align: top; }}
            .doc-table th {{ background-color: #f1f5f9; color: #1e293b; width: 32%; font-weight: bold; font-size: 12px; text-transform: uppercase; }}
            
            .signature-box {{ margin-top: 40px; display: flex; justify-content: space-between; page-break-inside: avoid; }}
            .sig-block {{ width: 45%; border: 1px dashed #94a3b8; height: 110px; padding: 10px; border-radius: 4px; background: #fafafa; }}
            .sig-title {{ font-size: 11px; font-weight: bold; color: #475569; text-transform: uppercase; border-bottom: 1px solid #cbd5e1; padding-bottom: 5px; }}
            
            .no-print {{ margin-bottom: 15px; text-align: right; }}
            .print-btn {{ background-color: #0284c7; color: white; border: none; padding: 12px 24px; font-size: 14px; font-weight: bold; border-radius: 6px; cursor: pointer; box-shadow: 0 2px 4px rgba(0,0,0,0.15); }}
            @media print {{ .no-print {{ display: none !important; }} body {{ margin: 0; }} }}
        </style>
    </head>
    <body>
        <div class="no-print">
            <button class="print-btn" onclick="window.print()">🖨️ Imprimer / Enregistrer en PDF</button>
        </div>
        <div class="header-box">
            <div class="header-title">FICHE DE SUIVI ET DE CONTRÔLE DE CHANTIER</div>
            <div class="header-sub">Chantier : {chantier_name}</div>
        </div>

        <table class="doc-table">
            <tr>
                <th>🗓️ Date de Réalisation</th>
                <td><b>{date_val}</b></td>
            </tr>
            <tr>
                <th>📌 Nature des Travaux</th>
                <td><b style="color: #1e3a8a; font-size: 14px;">{nature_val}</b></td>
            </tr>
            <tr>
                <th>🧱 Partie d'Ouvrage</th>
                <td><b style="font-size: 14px;">{partie_val}</b></td>
            </tr>
            <tr>
                <th>📍 Situation / PK</th>
                <td><b>{situ_val}</b></td>
            </tr>
            <tr>
                <th>🚜 Activité Réalisée</th>
                <td>{act_val}</td>
            </tr>
            <tr>
                <th>🧪 Essai / Contrôle Réalisé</th>
                <td><b>{essai_val}</b></td>
            </tr>
            <tr>
                <th>📑 Référence de Procédure</th>
                <td><code>{proc_val}</code></td>
            </tr>
            <tr>
                <th>📎 Pièces Jointes / Contrôles</th>
                <td>{pieces_val}</td>
            </tr>
        </table>

        <div class="signature-box">
            <div class="sig-block">
                <div class="sig-title">Visa Chef de Chantier / Conducteur :</div>
            </div>
            <div class="sig-block">
                <div class="sig-title">Visa Contrôle Qualité / Laboratoire :</div>
            </div>
        </div>
    </body>
    </html>
    """

# ==========================================
# REMPLACEMENT EN MODÈLE WORD
# ==========================================
def fill_word_template(template_file, replacements):
    doc = Document(template_file)
    
    def replace_in_paragraph(paragraph):
        for key, val in replacements.items():
            if key in paragraph.text:
                val_str = str(val) if pd.notna(val) and val is not None else ""
                paragraph.text = paragraph.text.replace(key, val_str)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# FONCTIONS UTILITAIRES EXCEL
# ==========================================
def save_to_excel_with_formatting(df_to_save, sheet_name):
    try:
        if not os.path.exists(FILE_PATH):
            with pd.ExcelWriter(FILE_PATH, engine="openpyxl") as writer:
                df_to_save.to_excel(writer, sheet_name=sheet_name, index=False)

        with pd.ExcelWriter(FILE_PATH, engine="openpyxl", mode="a", if_sheet_exists="replace") as writer:
            df_to_save.to_excel(writer, sheet_name=sheet_name, index=False)
            worksheet = writer.sheets[sheet_name]
            
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    val_str = str(cell.value or "")
                    if len(val_str) > max_length:
                        max_length = len(val_str)
                worksheet.column_dimensions[column_letter].width = min(max_length + 4, 60)
            
            num_rows = max(len(df_to_save) + 1, 2)
            num_cols = len(df_to_save.columns)
            
            if num_cols > 0:
                end_col_letter = get_column_letter(num_cols)
                table_ref = f"A1:{end_col_letter}{num_rows}"
                clean_name = re.sub(r'\W+', '_', sheet_name)
                table_name = f"Tableau_{clean_name}"
                
                worksheet._tables.clear()
                tab = Table(displayName=table_name, ref=table_ref)
                style = TableStyleInfo(name="TableStyleMedium3", showRowStripes=True)
                tab.tableStyleInfo = style
                worksheet.add_table(tab)

        return True, "✅ Fichier mis à jour !"
    except PermissionError:
        return False, "❌ ERREUR : Le fichier Excel est OUVERT."
    except Exception as e:
        return False, f"❌ Erreur : {e}"

def init_excel():
    if not os.path.exists(FILE_PATH):
        df_empty = pd.DataFrame(columns=COLUMNS_TEMPLATE)
        save_to_excel_with_formatting(df_empty, "Chantier Principal")

def get_sheet_names():
    if os.path.exists(FILE_PATH):
        try:
            return pd.ExcelFile(FILE_PATH).sheet_names
        except:
            return ["Chantier Principal"]
    return ["Chantier Principal"]

def load_data(sheet_name):
    try:
        df = pd.read_excel(FILE_PATH, sheet_name=sheet_name)
        if 'DATE' in df.columns:
            df['DATE'] = pd.to_datetime(df['DATE'], errors='coerce').dt.strftime('%d/%m/%Y')
        return df
    except Exception:
        return pd.DataFrame(columns=COLUMNS_TEMPLATE)

init_excel()
chantiers_disponibles = get_sheet_names()

# BARRE LATÉRALE
with st.sidebar:
    st.header("🏢 Gestion des Chantiers")
    chantier_actif = st.selectbox("📌 Chantier Actif :", options=chantiers_disponibles)
    st.markdown("---")
    st.subheader("➕ Ajouter un Chantier")
    nouveau_chantier = st.text_input("Nom du nouveau chantier :", placeholder="Ex: Tronçon B, Viaduc ...")
    if st.button("➕ Créer la feuille", use_container_width=True):
        if nouveau_chantier.strip() and nouveau_chantier.strip() not in chantiers_disponibles:
            save_to_excel_with_formatting(pd.DataFrame(columns=COLUMNS_TEMPLATE), nouveau_chantier.strip())
            st.rerun()

if not chantier_actif:
    st.stop()

# INTERFACE PRINCIPALE
st.title(f"🏗️ Suivi de Chantier : **{chantier_actif}**")
st.markdown("---")

df = load_data(chantier_actif)
tab1, tab2, tab3 = st.tabs(["📝 Nouvelle Saisie", "📊 Données & Filtres Excel", "📄 Générer Documents & Modèles"])

# ONGLET 1 : SAISIE
with tab1:
    st.subheader(f"📝 Saisie pour {chantier_actif}")
    col1, col2 = st.columns(2)
    with col1:
        date_saisie = st.date_input("🗓️ Date", value=datetime.today(), format="DD/MM/YYYY")
        nature_selectionnee = st.selectbox("📌 TITRE DE LA NATURE DES TRAVAUX", options=list(LIAISONS.keys()))
        info_liaison = LIAISONS.get(nature_selectionnee, {"procedure": "", "pieces": ""})
        partie_ouvrage = st.text_input("🧱 PARTIE D'OUVRAGE", placeholder="Ex: BRETELLE - A, BRANCHE 2...")
        situation = st.text_input("📍 SITUATION / PK" , placeholder="Ex: DU PK 0+020 AU PK 0+300")
    with col2:
        activite = st.text_area("🚜 ACTIVITÉ RÉALISÉE", height=80 ,placeholder="Ex: 1ér couche, 2éme couche ..." )
        essai = st.selectbox("🧪 ÉSSAI / CONTRÔLE RÉALISÉE", options=["Aucun", "ESSAI À LA PLAQUE", "DENSITÉ", "ESSAI À LA PLAQUE + DENSITÉ", "TENEUR EN EAU", "IDENTIFICATION DES MATERIAUX", "PRELEVEMENT AVANT COMPACTAGE", "PRELEVEMENT APRES COMPACTAGE", "PRELEVEMENT"])
        procedure = st.text_input("📑 RÉFÉRENCE DE PROCÉDURE", value=info_liaison["procedure"])
        pieces_jointes = st.text_area("📎 PIÈCES JOINTES", value=info_liaison["pieces"], height=120)

    if st.button(f"💾 Enregistrer dans {chantier_actif}", type="primary", use_container_width=True):
        new_entry = {
            "DATE": date_saisie.strftime('%d/%m/%Y'),
            "TITRE DE LA NATURE DES TRAVAUX": nature_selectionnee,
            COL_PARTIE: partie_ouvrage,
            "SITUATION": situation,
            "ACTIVITÉ RÉALISÉE": activite,
            "ÉSSAI/ CONTRÔLE RÉALISÉE": None if essai == "Aucun" else essai,
            "RÉFÉRENCE DE PROCÉDURE": procedure,
            "PIÈCES JOINTES": pieces_jointes
        }
        df_updated = pd.concat([load_data(chantier_actif), pd.DataFrame([new_entry])], ignore_index=True)
        save_to_excel_with_formatting(df_updated, chantier_actif)
        st.success("Enregistré !")
        st.rerun()

# ONGLET 2 : DONNÉES ET TABLEAU
with tab2:
    st.subheader(f"📊 Données & Filtres : {chantier_actif}")
    if not df.empty:
        filtered_df = df.copy()
        with st.expander("🔍 **Filtres & Recherche** *(cliquer pour ouvrir)*", expanded=False):
            c_search, c_annee, c_mois, c_jour = st.columns([3, 1, 1, 1])
            with c_search:
                search_query = st.text_input("⚡ Recherche globale")
                if search_query:
                    mask = filtered_df.apply(lambda col: col.astype(str).str.contains(search_query, case=False, na=False, regex=False)).any(axis=1)
                    filtered_df = filtered_df[mask]

            if 'DATE' in df.columns:
                temp_date = pd.to_datetime(filtered_df['DATE'], format='%d/%m/%Y', errors='coerce')
                with c_annee:
                    annees_sel = st.multiselect("Année", options=sorted(temp_date.dt.year.dropna().unique().astype(int).tolist()))
                with c_mois:
                    mois_sel = st.multiselect("Mois", options=[str(m).zfill(2) for m in sorted(temp_date.dt.month.dropna().unique().astype(int).tolist())])
                with c_jour:
                    jours_sel = st.multiselect("Jour", options=[str(j).zfill(2) for j in sorted(temp_date.dt.day.dropna().unique().astype(int).tolist())])

                if annees_sel: filtered_df = filtered_df[temp_date.dt.year.isin(annees_sel)]
                if mois_sel: filtered_df = filtered_df[temp_date.dt.month.isin([int(m) for m in mois_sel])]
                if jours_sel: filtered_df = filtered_df[temp_date.dt.day.isin([int(j) for j in jours_sel])]

            st.divider()
            filter_cols = st.columns(3)
            col_idx = 0
            for col_name in df.columns:
                if col_name == 'DATE': continue
                unique_vals = sorted([str(v) for v in df[col_name].dropna().unique()])
                with filter_cols[col_idx % 3]:
                    selected = st.multiselect(f"📌 {col_name}", options=unique_vals, key=f"filter_{col_name}")
                    if selected: filtered_df = filtered_df[filtered_df[col_name].astype(str).isin(selected)]
                col_idx += 1

        st.caption(f"📊 **{len(filtered_df)}** ligne(s) affichée(s) sur **{len(df)}**")
        edited_df = st.data_editor(filtered_df, num_rows="dynamic", use_container_width=True, height=350, key="editor")
        
        col_action1, col_action2 = st.columns(2)
        with col_action1:
            if st.button("💾 Enregistrer les modifications Excel", type="primary", use_container_width=True):
                df.update(edited_df)
                save_to_excel_with_formatting(df, chantier_actif)
                st.success("Modifications enregistrées !")
                st.rerun()

        with col_action2:
            if st.button("🖨️ Imprimer le tableau complet (PDF / Papier)", use_container_width=True):
                st.session_state["show_print_view"] = True
                st.session_state["show_single_doc_view"] = False

        if st.session_state.get("show_print_view", False):
            st.markdown("---")
            st.subheader("🖨️ Vue d'Impression du Tableau Général")
            printable_html = generate_printable_html(edited_df, chantier_actif)
            c_dl, c_close = st.columns([3, 1])
            with c_dl:
                st.download_button(
                    label="📥 Télécharger l'Impression Tableau (HTML)",
                    data=printable_html,
                    file_name=f"Rapport_{chantier_actif}.html",
                    mime="text/html",
                    use_container_width=True
                )
            with c_close:
                if st.button("❌ Masquer l'impression", use_container_width=True):
                    st.session_state["show_print_view"] = False
                    st.rerun()

            components.html(printable_html, height=500, scrolling=True)

# ONGLET 3 : GÉNÉRER DOCUMENTS & MODÈLES
with tab3:
    st.subheader("📄 Génération de Documents Individuels (PDF & Word)")
    
    if not df.empty:
        st.markdown("### 1. Sélectionner une ligne pour la fiche / document")
        options_lignes = {
            i: f"Ligne {i+1} | {df.loc[i, 'DATE']} | {df.loc[i, 'TITRE DE LA NATURE DES TRAVAUX']} | {df.loc[i, COL_PARTIE]}"
            for i in df.index
        }
        selected_index = st.selectbox("Choisir la ligne de travaux :", options=list(options_lignes.keys()), format_func=lambda x: options_lignes[x])
        row_selected = df.loc[selected_index]

        nat_txt = str(row_selected.get('TITRE DE LA NATURE DES TRAVAUX', 'NATURE')).strip()
        part_txt = str(row_selected.get(COL_PARTIE, 'PARTIE')).strip()
        doc_pdf_filename = f"{nat_txt}_{part_txt} _.pdf"

        st.markdown("---")
        
        doc_col1, doc_col2 = st.columns(2)

        # OPTION A : FICHE PDF INDIVIDUELLE (COMMANDE EXCEL)
        with doc_col1:
            st.markdown("### 🖨️ Option A : Générer le Document PDF")
            st.info("Génère la fiche individuelle de contrôle au format PDF avec le nom structuré.")
            
            if st.button("📄 Générer le document PDF", type="primary", use_container_width=True):
                st.session_state["show_single_doc_view"] = True
                st.session_state["pdf_success_msg"] = f"Le PDF a été généré avec succès : {doc_pdf_filename}"

        # OPTION B : FICHIER WORD MODÈLE (.DOCX)
        with doc_col2:
            st.markdown("### 📝 Option B : Remplir un Modèle Word (.docx)")
            uploaded_word = st.file_uploader("Importer votre fichier modèle Word", type=["docx"])

            if uploaded_word is not None:
                replacements = {
                    "{{DATE}}": str(row_selected.get('DATE', '')),
                    "{{NATURE}}": str(row_selected.get('TITRE DE LA NATURE DES TRAVAUX', '')),
                    "{{PARTIE}}": str(row_selected.get(COL_PARTIE, '')),
                    "{{SITUATION}}": str(row_selected.get('SITUATION', '')),
                    "{{ACTIVITE}}": str(row_selected.get('ACTIVITÉ RÉALISÉE', '')),
                    "{{ESSAI}}": str(row_selected.get('ÉSSAI/ CONTRÔLE RÉALISÉE', '')),
                    "{{REF}}": str(row_selected.get('RÉFÉRENCE DE PROCÉDURE', '')),
                    "{{PIECES}}": str(row_selected.get('PIÈCES JOINTES', ''))
                }
                
                doc_output_name = f"{nat_txt}_{part_txt}.docx"
                filled_doc = fill_word_template(uploaded_word, replacements)
                
                st.download_button(
                    label=f"📥 Télécharger le Document Word ({doc_output_name})",
                    data=filled_doc,
                    file_name=doc_output_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        # AFFICHAGE DU DOCUMENT PDF GENERÉ
        if st.session_state.get("show_single_doc_view", False):
            st.markdown("---")
            if st.session_state.get("pdf_success_msg"):
                st.success(f"ℹ️ **{st.session_state['pdf_success_msg']}**")
            
            single_doc_html = generate_single_document_html(row_selected, chantier_actif)
            
            cdl_doc, cclose_doc = st.columns([3, 1])
            with cdl_doc:
                st.download_button(
                    label=f"📥 Télécharger la Fiche HTML / PDF",
                    data=single_doc_html,
                    file_name=f"{nat_txt}_{part_txt}.html",
                    mime="text/html",
                    use_container_width=True
                )
            with cclose_doc:
                if st.button("❌ Masquer le document", use_container_width=True):
                    st.session_state["show_single_doc_view"] = False
                    st.rerun()

            components.html(single_doc_html, height=650, scrolling=True)

    else:
        st.info("Aucune donnée disponible dans le tableau.")